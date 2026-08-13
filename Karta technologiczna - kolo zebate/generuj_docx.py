#!/usr/bin/env python3
"""Generuje edytowalna w Wordzie karte technologiczna (.docx) kola zebatego.

Odwzorowuje formatke (naglowek + tabela operacji) z wypelniona kolumna
'Stanowisko' (obrabiarki z katalogu) oraz 'Pomoce warsztatowe' (przyrzady
pomiarowe + narzedzia skrawajace).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DST = "karta_technologiczna_uzupelniona.docx"

# ---------------------------------------------------------------- dane operacji
# (nr, stanowisko, opis operacji, pomoce warsztatowe)
OPERACJE = [
    ("5", "BFO\nPrzecinarka taśmowa automatyczna BEE 250",
     "Odcięcie materiału wyjściowego z pręta walcowanego na określoną długość "
     "z uwzględnieniem naddatku na obróbkę powierzchni czołowych.",
     "Suwmiarka"),
    ("10", "JAFO\nFrezarka uniwersalna wspornikowa FWF-32J2",
     "Frezowanie zgrubne powierzchni czołowych.",
     "Imadło maszynowe\nSuwmiarka\nCzujnik zegarowy"),
    ("15", "CHOFUM\nTokarka uchwytowa TZC-32N1",
     "Toczenie zgrubne i średniodokładne powierzchni walcowej.",
     "Uchwyt 3-szczękowy\nmikrometr zewnętrzny\nsuwmiarka"),
    ("20", "CHOFUM\nTokarka uchwytowa TZC-32N1",
     "Wiercenie i powiercanie oraz frezowanie średnio dokładne i wykańczające "
     "otworu piasty.",
     "Imadło\nSuwmiarka\nUchwyt trójszczękowy"),
    ("25", "Dłutownica",
     "Dłutowanie zgrubne i średnio dokładne rowka wpustowego.",
     "Przyrząd podziałowy\nsuwmiarka głębokościowa\npłytki wzorcowe"),
    ("30", "Frezarka obwiedniowa do uzębień",
     "Frezowanie obwiedniowe zgrubne i średnio dokładne uzębienia.",
     "Trzpień montażowy\nmikrometr talerzykowy"),
    ("35", "JAFO\nFrezarka uniwersalna FWF-32J2\n(podzielnica)",
     "Wiercenie otworów poprzecznych, frezowanie zgrubne powierzchni "
     "odciążających oraz wykonanie faz frezem fazującym.",
     "Imadło maszynowe\nUchwyt podziałowy\nsuwmiarka"),
    ("40", "Stanowisko hartowania płomieniowego",
     "Hartowanie ogniowe.",
     "Termometr bezdotykowy"),
    ("45", "Piec do odpuszczania",
     "Niskie odpuszczanie po procesie hartowania w celu likwidacji naprężeń "
     "własnych i poprawy udarności rdzenia.",
     "termoogniwa kontrolne z rejestratorem profili temperatur"),
    ("50", "Szlifierka do uzębień",
     "Ostateczna obróbka wykańczająca profili zębów po hartowaniu poprzez szlifowanie.",
     "trzpień precyzyjny\nmikrometr talerzykowy"),
    ("55", "Myjnia przemysłowa",
     "Dokładne mycie i czyszczenie.",
     "Kosze transportowe\nalkaliczny środek myjący\npistolet ze sprężonym powietrzem"),
    ("60", "IOS\nWspółrzędnościowa maszyna pomiarowa MP 700E",
     "Ostateczna kontrola wymiarów gotowego wyrobu.",
     "Współrzędnościowa maszyna pomiarowa\newolwentomierz do kontroli odchyłek profilu i linii zęba"),
]


def set_cell_border(cell, **kw):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tcBorders.append(el)


def style_cell(cell, text, size=8, bold=False, align='left', valign='center'):
    cell.vertical_alignment = (WD_ALIGN_VERTICAL.CENTER if valign == 'center'
                               else WD_ALIGN_VERTICAL.TOP)
    cell.text = ""
    lines = str(text).split("\n")
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                       'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = "Arial"
    set_cell_border(cell)


def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Cm(1.3))

    # ---- naglowek (metryczka) ----
    head = doc.add_table(rows=3, cols=5)
    head.style = 'Table Grid'
    head.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdata = [
        ["Karta technologiczna\n(plan operacyjny)", "Wyrób", "Nazwa części\nKoło zębate",
         "Nr rys. części", "Znak"],
        ["Gat. i stan mat.", "Postać i wymiary półfabrykatu [mm]", "Sztuk / wyrób",
         "kg/1 szt. netto", "Sztuk na zlecenie, partię"],
        ["Staliwo", "192x52", "Norma mat. kg/1 szt.", "Materiał kg / zlecenie, partię", ""],
    ]
    for r, row in enumerate(hdata):
        for c, val in enumerate(row):
            big = (r == 0 and c == 0)
            style_cell(head.cell(r, c), val, size=11 if big else 8,
                       bold=big, align='center' if big else 'left', valign='center')
    set_col_widths(head, [Cm(4.6), Cm(4.5), Cm(3.6), Cm(3.0), Cm(2.6)])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- tabela operacji ----
    cols = ["Nr op", "Stanowisko", "Opis operacji", "Pomoce warsztatowe",
            "Kat.r dod.", "Tpz / tj", "T"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, name in enumerate(cols):
        style_cell(tbl.rows[0].cells[c], name, size=8, bold=True, align='center')

    for nr, stan, opis, pom in OPERACJE:
        cells = tbl.add_row().cells
        style_cell(cells[0], nr, size=8, bold=True, align='center')
        style_cell(cells[1], stan, size=7.5, align='center')
        style_cell(cells[2], opis, size=8, align='left')
        style_cell(cells[3], pom, size=7.5, align='left')
        style_cell(cells[4], "", size=8)
        style_cell(cells[5], "", size=8)
        style_cell(cells[6], "", size=8)

    widths = [Cm(1.0), Cm(3.4), Cm(6.1), Cm(4.3), Cm(1.3), Cm(1.3), Cm(0.9)]
    set_col_widths(tbl, widths)

    # ---- stopka ----
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    foot = doc.add_table(rows=1, cols=3)
    foot.style = 'Table Grid'
    foot.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, val in enumerate(["Opracował: Jakub Wydra", "Sprawdził: Kacper Tokarski",
                             "Zatwierdził: Maciej Tkacz"]):
        style_cell(foot.cell(0, c), val, size=9, align='left')
    set_col_widths(foot, [Cm(6.2), Cm(6.2), Cm(6.0)])

    doc.save(DST)
    print("Zapisano:", DST)


if __name__ == "__main__":
    main()
