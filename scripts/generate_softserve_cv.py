from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


OUT_DIR = Path(__file__).resolve().parents[1] / "artifacts" / "softserve-cv"
FONT_REGULAR = "DejaVuSans"
FONT_BOLD = "DejaVuSans-Bold"

pdfmetrics.registerFont(TTFont(FONT_REGULAR, "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
pdfmetrics.registerFont(TTFont(FONT_BOLD, "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))


PROFILE_EN = (
    "Mechanical Engineering student at Cracow University of Technology and certified Software "
    "Technician combining CAD/mechanical design with Python/C++ programming, robotic simulation "
    "and hands-on prototyping. Commercial experience creating robotics simulations and digital "
    "twins in Visual Components for automotive production environments. Built an end-to-end "
    "SCARA robotic arm prototype covering CAD, FDM printing, actuator/driver selection, "
    "electronics, control software and hardware-software integration. Interested in automation, "
    "mobile/industrial robotics, drones and simulation-driven development."
)

PROFILE_PL = (
    "Student Mechaniki i Budowy Maszyn na Politechnice Krakowskiej oraz Technik Programista, "
    "łączący projektowanie CAD/mechaniczne z programowaniem w Pythonie/C++, symulacjami robotycznymi "
    "i praktycznym prototypowaniem. Posiadam komercyjne doświadczenie w tworzeniu symulacji "
    "robotycznych i cyfrowych bliźniaków w Visual Components dla środowisk produkcyjnych branży "
    "Automotive. Zrealizowałem end-to-end prototyp robota SCARA obejmujący CAD, druk FDM, dobór "
    "napędów i sterowników, elektronikę, oprogramowanie sterujące oraz integrację hardware-software. "
    "Interesuję się automatyką, robotyką mobilną/przemysłową, dronami i rozwojem systemów z użyciem symulacji."
)


CV_EN = {
    "filename": "Maciej_Tkacz_CV_SoftServe_EN",
    "name": "Maciej Tkacz",
    "title": "Junior Robotics / Simulation Engineer",
    "contact": [
        "Email: maciek01110@gmail.com",
        "Phone: +48 881 912 125",
        "Location: Krakow, Poland",
        "English: C1",
        "Driving licence: category B",
    ],
    "profile_title": "Professional Profile",
    "profile": PROFILE_EN,
    "sections": [
        {
            "title": "Key Skills",
            "items": [
                "Robotics simulation & digital twins: Visual Components, robot workcells, robot paths, process logic, collision checks, cycle flows and layout optimization.",
                "Programming: Python, C++, Arduino/C, JavaScript basics; automation scripts, calculation scripts, hardware-control experiments and technical troubleshooting.",
                "Robotics & control: kinematics fundamentals, motion basics, sensing concepts, control systems, dynamic systems modelling, automation, mechatronics and analytical mechanics coursework.",
                "Simulation exposure: MATLAB at university; ROS self-study in personal robotics projects; strong willingness to learn Gazebo, Isaac Sim and similar tools.",
                "CAD & prototyping: SolidWorks, Autodesk Inventor, Fusion 360, AutoCAD; rapid prototyping with FDM, SLA and MJF; design for additive manufacturing and reverse engineering.",
                "Hardware-software integration: stepper motors, motor drivers, microcontroller/single-board platforms, wiring, testing and iterative debugging of electromechanical prototypes.",
                "Teamwork: technical documentation, communication with engineering teams and collaborative problem solving in interdisciplinary environments.",
            ],
        },
        {
            "title": "Selected Robotics Project",
            "subtitle": "SCARA Robotic Arm Prototype | Personal R&D Project",
            "items": [
                "Designed and built an end-to-end SCARA robotic arm prototype, covering mechanical concept, CAD modelling, kinematic layout, actuator/driver selection, electronics and wiring.",
                "Optimized components for 3D printing, including PET-G and carbon-fibre reinforced materials, with focus on stiffness, assembly constraints and rapid iteration.",
                "Developed Python/C++ control software for motion and hardware-interaction experiments, applying forward/inverse kinematics and control-system fundamentals.",
                "Integrated mechanical, electronic and software subsystems through iterative assembly, testing and debugging, gaining practical hardware-software integration experience.",
            ],
        },
        {
            "title": "Professional Experience",
            "jobs": [
                {
                    "role": "Application Engineer",
                    "company": "AIAutomation",
                    "date": "Aug 2025 - May 2026",
                    "items": [
                        "Created robotics simulations and digital twins of production workcells in Visual Components for automotive clients.",
                        "Developed robot logic, motion sequences, robot paths, collision checks, cycle flows and virtual process validation.",
                        "Optimized workcell layouts and 3D geometry for simulation needs, using SolidWorks and Autodesk Inventor.",
                        "Analyzed customer documentation, engineering standards and technical specifications to support solution concepts aligned with production requirements.",
                        "Collaborated with engineers and stakeholders, communicating simulation assumptions, constraints and improvement proposals.",
                    ],
                },
                {
                    "role": "3D Printing and CAD Design Specialist",
                    "company": "Cubic Inch Additive Manufacturing, Piaseczno",
                    "date": "Jun 2023 - Aug 2023",
                    "items": [
                        "Operated and serviced FDM, MJF and SLA 3D printers, supervising process parameters, post-processing and quality control.",
                        "Designed and optimized CAD models in Fusion 360 and Autodesk Inventor for additive manufacturing and rapid prototyping.",
                        "Supported implementation testing for a new SLA technology, documenting progress and technical observations.",
                        "Coordinated production tasks in a 10-person project team, balancing quality, manufacturability and delivery constraints.",
                    ],
                },
                {
                    "role": "Robotics and 3D Printing Intern",
                    "company": "ASTOR Robotics Center, Krakow",
                    "date": "May 2022",
                    "items": [
                        "Assembled mechanical equipment and 3D-printed components for Kawasaki robots and Astorino educational robot platforms.",
                        "Programmed Kawasaki industrial robots using teach pendant workflows and created basic robot motion programs.",
                        "Tested robot and workstation operation, gaining practical exposure to robot setup, safety and hardware integration.",
                    ],
                },
                {
                    "role": "Web Application Development Intern",
                    "company": "Souczek Design Studio Reklamy i Druku, Kielce",
                    "date": "Jul 2021",
                    "items": [
                        "Solved technical problems in JavaScript, including implementation of an interactive order form.",
                        "Tested and deployed web functionality in a team environment according to customer guidelines.",
                    ],
                },
            ],
        },
        {
            "title": "Education",
            "items": [
                "Cracow University of Technology - Mechanical Engineering, Engineer's degree in progress (Oct 2024 - present). Relevant coursework: dynamic systems modelling, automation/control, analytical mechanics and mechatronics.",
                "PKMechPower Student Research Group, Mechanical Section - CAD and 3D printing tasks, including driver seat and brake-system components; mass and strength-oriented design optimization.",
                "Zespol Szkol Informatycznych im. gen. Jozefa Hauke-Bosaka, Kielce - Software Technician (Sep 2019 - Apr 2024).",
            ],
        },
        {
            "title": "Certificates",
            "items": [
                "Kawasaki robot operation and programming - integrator course, ASTOR Robotics Center (May 2022).",
                "Python programming courses and self-directed Python/C++ development in robotics and automation projects.",
            ],
        },
    ],
    "sidebar": {
        "Contact": [
            "maciek01110@gmail.com",
            "+48 881 912 125",
            "Krakow, Poland",
        ],
        "Languages": ["Polish - native", "English - C1"],
        "Technical Stack": [
            "Python",
            "C++ / Arduino C",
            "Visual Components",
            "MATLAB basics",
            "ROS self-study",
            "SolidWorks",
            "Autodesk Inventor",
            "Fusion 360",
            "AutoCAD",
            "FDM / SLA / MJF",
            "Digital twins",
            "Hardware-software integration",
        ],
        "Robotics Keywords": [
            "Robotic manipulators",
            "SCARA prototype",
            "Robot paths",
            "Collision checking",
            "Control systems",
            "Dynamic systems",
            "Rapid prototyping",
        ],
    },
    "consent": "I hereby consent to the processing of my personal data for the purpose of the current recruitment process.",
}


CV_PL = {
    "filename": "Maciej_Tkacz_CV_SoftServe_PL",
    "name": "Maciej Tkacz",
    "title": "Junior Robotics / Simulation Engineer",
    "contact": [
        "E-mail: maciek01110@gmail.com",
        "Telefon: +48 881 912 125",
        "Lokalizacja: Kraków",
        "Język angielski: C1",
        "Prawo jazdy: kat. B",
    ],
    "profile_title": "Profil zawodowy",
    "profile": PROFILE_PL,
    "sections": [
        {
            "title": "Kluczowe umiejętności",
            "items": [
                "Symulacje robotyczne i digital twins: Visual Components, gniazda zrobotyzowane, ścieżki ruchu robotów, logika procesu, wykrywanie kolizji, cykle pracy i optymalizacja layoutu.",
                "Programowanie: Python, C++, Arduino/C, podstawy JavaScript; skrypty automatyzujące, obliczeniowe, do sterowania hardware'em i rozwiązywania problemów technicznych.",
                "Robotyka i sterowanie: podstawy kinematyki, ruchu, sensoryki, układów sterowania, modelowania układów dynamicznych, automatyki, mechatroniki i mechaniki analitycznej.",
                "Narzędzia symulacyjne: MATLAB na studiach; ROS rozwijany samodzielnie w projektach robotycznych; gotowość do nauki Gazebo, Isaac Sim i podobnych środowisk.",
                "CAD i prototypowanie: SolidWorks, Autodesk Inventor, Fusion 360, AutoCAD; rapid prototyping w FDM, SLA i MJF; projektowanie pod technologie przyrostowe i inżynieria odwrotna.",
                "Integracja hardware-software: silniki krokowe, sterowniki silników, platformy mikrokontrolerowe/jednopłytkowe, okablowanie, testowanie i iteracyjne debugowanie prototypów elektromechanicznych.",
                "Praca zespołowa: dokumentacja techniczna, komunikacja z zespołami inżynieryjnymi i rozwiązywanie problemów w środowisku interdyscyplinarnym.",
            ],
        },
        {
            "title": "Wybrany projekt robotyczny",
            "subtitle": "Prototyp ramienia robota SCARA | Projekt własny R&D",
            "items": [
                "Zaprojektowanie i budowa end-to-end prototypu ramienia SCARA: koncepcja mechaniczna, modelowanie CAD, układ kinematyczny, dobór napędów/sterowników, elektronika i okablowanie.",
                "Optymalizacja części pod druk 3D, w tym PET-G i materiały wzmacniane włóknem węglowym, z naciskiem na sztywność, montaż i szybkie iteracje konstrukcji.",
                "Tworzenie oprogramowania sterującego w Pythonie/C++ do eksperymentów z ruchem i komunikacją z hardware'em, z wykorzystaniem podstaw kinematyki prostej/odwrotnej i sterowania.",
                "Integracja podsystemów mechanicznych, elektronicznych i software'owych poprzez iteracyjny montaż, testowanie i debugowanie.",
            ],
        },
        {
            "title": "Doświadczenie zawodowe",
            "jobs": [
                {
                    "role": "Application Engineer",
                    "company": "AIAutomation",
                    "date": "08.2025 - 05.2026",
                    "items": [
                        "Tworzenie symulacji robotycznych i cyfrowych bliźniaków gniazd produkcyjnych w Visual Components dla klientów z branży Automotive.",
                        "Przygotowywanie logiki pracy robotów, sekwencji ruchu, ścieżek, kontroli kolizji, cykli produkcyjnych i wirtualnej walidacji procesu.",
                        "Optymalizacja layoutów stanowisk oraz geometrii 3D na potrzeby symulacji z użyciem SolidWorks i Autodesk Inventor.",
                        "Analiza dokumentacji klienta, standardów inżynieryjnych i specyfikacji technicznych w celu wsparcia koncepcji rozwiązań zgodnych z wymaganiami produkcyjnymi.",
                        "Współpraca z inżynierami i interesariuszami oraz komunikowanie założeń, ograniczeń i propozycji usprawnień w symulacji.",
                    ],
                },
                {
                    "role": "Specjalista ds. Druku 3D i Projektowania CAD",
                    "company": "Cubic Inch Additive Manufacturing, Piaseczno",
                    "date": "06.2023 - 08.2023",
                    "items": [
                        "Obsługa i serwis drukarek 3D FDM, MJF i SLA: nadzór nad parametrami procesu, post-processing i kontrola jakości.",
                        "Projektowanie i optymalizacja modeli CAD w Fusion 360 i Autodesk Inventor pod technologie przyrostowe oraz rapid prototyping.",
                        "Wsparcie wdrożeniowe nowej technologii SLA poprzez testy, dokumentowanie postępów i obserwacji technicznych.",
                        "Koordynacja zadań produkcyjnych w 10-osobowym zespole projektowym z uwzględnieniem jakości, wytwarzalności i terminowości.",
                    ],
                },
                {
                    "role": "Praktykant ds. Robotyki i Druku 3D",
                    "company": "ASTOR Robotics Center, Kraków",
                    "date": "05.2022",
                    "items": [
                        "Montaż osprzętu mechanicznego i komponentów drukowanych 3D do robotów Kawasaki oraz edukacyjnych platform Astorino.",
                        "Programowanie robotów przemysłowych Kawasaki z wykorzystaniem teach pendanta oraz tworzenie podstawowych programów ruchu.",
                        "Testowanie działania robotów i stanowisk, zdobywanie praktycznej wiedzy z zakresu konfiguracji, bezpieczeństwa i integracji hardware'u.",
                    ],
                },
                {
                    "role": "Stażysta ds. Tworzenia Aplikacji Webowych",
                    "company": "Souczek Design Studio Reklamy i Druku, Kielce",
                    "date": "07.2021",
                    "items": [
                        "Rozwiązywanie problemów technicznych w JavaScript, w tym wdrożenie interaktywnego formularza zamówień.",
                        "Testowanie i wdrażanie funkcjonalności webowych w zespole zgodnie z wytycznymi klienta.",
                    ],
                },
            ],
        },
        {
            "title": "Wykształcenie",
            "items": [
                "Politechnika Krakowska im. Tadeusza Kościuszki - Mechanika i Budowa Maszyn, studia inżynierskie w toku (10.2024 - obecnie). Istotne obszary: modelowanie układów dynamicznych, automatyka/sterowanie, mechanika analityczna i mechatronika.",
                "Koło Naukowe PKMechPower, Sekcja Mechaniczna - zadania konstrukcyjne w CAD i druku 3D, m.in. projekt fotela kierowcy i podzespołów układu hamulcowego; optymalizacja pod kątem masy i wytrzymałości.",
                "Zespół Szkół Informatycznych im. gen. Józefa Hauke-Bosaka w Kielcach - Technik programista (09.2019 - 04.2024).",
            ],
        },
        {
            "title": "Certyfikaty",
            "items": [
                "Obsługa i programowanie robotów Kawasaki - kurs dla integratorów, ASTOR Robotics Center (05.2022).",
                "Kursy programowania w Pythonie oraz samodzielny rozwój w Python/C++ w projektach robotycznych i automatyzacyjnych.",
            ],
        },
    ],
    "sidebar": {
        "Kontakt": [
            "maciek01110@gmail.com",
            "+48 881 912 125",
            "Kraków",
        ],
        "Języki": ["Polski - ojczysty", "Angielski - C1"],
        "Stack techniczny": [
            "Python",
            "C++ / Arduino C",
            "Visual Components",
            "Podstawy MATLAB",
            "ROS - samodzielna nauka",
            "SolidWorks",
            "Autodesk Inventor",
            "Fusion 360",
            "AutoCAD",
            "FDM / SLA / MJF",
            "Digital twins",
            "Integracja hardware-software",
        ],
        "Słowa kluczowe": [
            "Manipulatory robotyczne",
            "Prototyp SCARA",
            "Ścieżki robotów",
            "Kontrola kolizji",
            "Układy sterowania",
            "Układy dynamiczne",
            "Rapid prototyping",
        ],
    },
    "consent": "Wyrażam zgodę na przetwarzanie moich danych osobowych w celu prowadzenia obecnego postępowania rekrutacyjnego.",
}


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)


def add_doc_heading(paragraph, text: str, size: int, bold: bool = True, color: str = "1F4E79") -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_doc_bullets(container, items: Iterable[str], style: str = "List Bullet") -> None:
    for item in items:
        p = container.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.35)
        p.add_run(item)


def add_section_doc(doc_or_cell, title: str) -> None:
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    add_doc_heading(p, title.upper(), 10)


def build_ats_docx(cv: dict) -> Path:
    doc = Document()
    set_doc_defaults(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_doc_heading(name, cv["name"], 20)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(cv["title"]).bold = True
    contact = doc.add_paragraph(" | ".join(cv["contact"]))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.runs[0].font.size = Pt(8)

    add_section_doc(doc, cv["profile_title"])
    doc.add_paragraph(cv["profile"])

    for section in cv["sections"]:
        add_section_doc(doc, section["title"])
        if "subtitle" in section:
            p = doc.add_paragraph()
            p.add_run(section["subtitle"]).bold = True
        if "items" in section:
            add_doc_bullets(doc, section["items"])
        if "jobs" in section:
            for job in section["jobs"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                r = p.add_run(job["role"])
                r.bold = True
                p.add_run(f" | {job['company']} | {job['date']}")
                add_doc_bullets(doc, job["items"])

    consent = doc.add_paragraph(cv["consent"])
    consent.runs[0].font.size = Pt(7)

    path = OUT_DIR / f"{cv['filename']}_ATS.docx"
    doc.save(path)
    return path


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_visual_docx(cv: dict) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)

    header = doc.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcell = header.cell(0, 0)
    shade_cell(hcell, "F2F6FA")
    p = hcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_doc_heading(p, cv["name"], 20)
    p2 = hcell.add_paragraph(cv["title"])
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.7)
    table.columns[1].width = Inches(2.0)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade_cell(right, "F2F6FA")

    add_section_doc(left, cv["profile_title"])
    left.add_paragraph(cv["profile"])

    for section_data in cv["sections"]:
        if section_data["title"] in ("Key Skills", "Kluczowe umiejętności"):
            continue
        add_section_doc(left, section_data["title"])
        if "subtitle" in section_data:
            p = left.add_paragraph()
            p.add_run(section_data["subtitle"]).bold = True
        if "items" in section_data:
            add_doc_bullets(left, section_data["items"])
        if "jobs" in section_data:
            for job in section_data["jobs"]:
                p = left.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.add_run(job["role"]).bold = True
                p.add_run(f" | {job['company']} | {job['date']}")
                add_doc_bullets(left, job["items"])

    for heading, items in cv["sidebar"].items():
        add_section_doc(right, heading)
        add_doc_bullets(right, items)

    add_section_doc(right, "Consent" if cv is CV_EN else "Zgoda")
    c = right.add_paragraph(cv["consent"])
    c.runs[0].font.size = Pt(7)

    path = OUT_DIR / f"{cv['filename']}_Visual.docx"
    doc.save(path)
    return path


def split_text(text: str, font: str, size: int, width: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if stringWidth(candidate, font, size) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


class PdfWriter:
    def __init__(self, path: Path):
        self.path = path
        self.c = canvas.Canvas(str(path), pagesize=A4)
        self.width, self.height = A4
        self.margin = 1.45 * cm
        self.y = self.height - self.margin

    def ensure(self, needed: float) -> None:
        if self.y - needed < self.margin:
            self.c.showPage()
            self.y = self.height - self.margin

    def text(self, text: str, x: float, width: float, size: int = 9, font: str = FONT_REGULAR, leading: float = 11, color=colors.black):
        self.c.setFont(font, size)
        self.c.setFillColor(color)
        lines = split_text(text, font, size, width)
        self.ensure(len(lines) * leading + 2)
        for line in lines:
            self.c.drawString(x, self.y, line)
            self.y -= leading
        return len(lines)

    def heading(self, text: str, x: float, width: float):
        self.ensure(22)
        self.y -= 5
        self.c.setFillColor(colors.HexColor("#1F4E79"))
        self.c.setFont(FONT_BOLD, 10)
        self.c.drawString(x, self.y, text.upper())
        self.y -= 4
        self.c.setStrokeColor(colors.HexColor("#1F4E79"))
        self.c.line(x, self.y, x + width, self.y)
        self.y -= 10

    def bullet(self, text: str, x: float, width: float):
        bullet_width = 10
        lines = split_text(text, FONT_REGULAR, 8.6, width - bullet_width)
        self.ensure(len(lines) * 10 + 2)
        self.c.setFillColor(colors.black)
        self.c.setFont(FONT_REGULAR, 8.6)
        self.c.drawString(x, self.y, "-")
        self.c.drawString(x + bullet_width, self.y, lines[0])
        self.y -= 10
        for line in lines[1:]:
            self.c.drawString(x + bullet_width, self.y, line)
            self.y -= 10
        self.y -= 1

    def save(self):
        self.c.save()


def build_ats_pdf(cv: dict) -> Path:
    path = OUT_DIR / f"{cv['filename']}_ATS.pdf"
    pdf = PdfWriter(path)
    x = pdf.margin
    width = pdf.width - 2 * pdf.margin

    pdf.c.setFillColor(colors.HexColor("#1F4E79"))
    pdf.c.setFont(FONT_BOLD, 20)
    pdf.c.drawCentredString(pdf.width / 2, pdf.y, cv["name"])
    pdf.y -= 18
    pdf.c.setFillColor(colors.black)
    pdf.c.setFont(FONT_BOLD, 11)
    pdf.c.drawCentredString(pdf.width / 2, pdf.y, cv["title"])
    pdf.y -= 15
    pdf.c.setFont(FONT_REGULAR, 8)
    pdf.c.drawCentredString(pdf.width / 2, pdf.y, " | ".join(cv["contact"]))
    pdf.y -= 12

    pdf.heading(cv["profile_title"], x, width)
    pdf.text(cv["profile"], x, width, size=8.8, leading=10.5)

    for section in cv["sections"]:
        pdf.heading(section["title"], x, width)
        if "subtitle" in section:
            pdf.text(section["subtitle"], x, width, size=9, font=FONT_BOLD, leading=11)
        for item in section.get("items", []):
            pdf.bullet(item, x, width)
        for job in section.get("jobs", []):
            pdf.text(f"{job['role']} | {job['company']} | {job['date']}", x, width, size=9, font=FONT_BOLD, leading=11)
            for item in job["items"]:
                pdf.bullet(item, x, width)

    pdf.y -= 5
    pdf.text(cv["consent"], x, width, size=6.8, leading=8)
    pdf.save()
    return path


def draw_sidebar(c: canvas.Canvas, cv: dict, x: float, y: float, width: float):
    c.setFillColor(colors.HexColor("#F2F6FA"))
    c.rect(x - 0.25 * cm, 0, width + 0.5 * cm, A4[1], stroke=0, fill=1)
    current_y = y
    for heading, items in cv["sidebar"].items():
        c.setFillColor(colors.HexColor("#1F4E79"))
        c.setFont(FONT_BOLD, 9)
        c.drawString(x, current_y, heading.upper())
        current_y -= 10
        c.setStrokeColor(colors.HexColor("#1F4E79"))
        c.line(x, current_y, x + width, current_y)
        current_y -= 10
        c.setFillColor(colors.black)
        c.setFont(FONT_REGULAR, 7.6)
        for item in items:
            lines = split_text(item, FONT_REGULAR, 7.6, width - 7)
            c.drawString(x, current_y, "-")
            c.drawString(x + 7, current_y, lines[0])
            current_y -= 9
            for line in lines[1:]:
                c.drawString(x + 7, current_y, line)
                current_y -= 9
            current_y -= 1
        current_y -= 7


def build_visual_pdf(cv: dict) -> Path:
    path = OUT_DIR / f"{cv['filename']}_Visual.pdf"
    c = canvas.Canvas(str(path), pagesize=A4)
    page_w, page_h = A4
    margin = 1.15 * cm
    sidebar_w = 5.15 * cm
    gap = 0.55 * cm
    main_x = margin
    sidebar_x = page_w - margin - sidebar_w
    main_w = sidebar_x - gap - main_x

    draw_sidebar(c, cv, sidebar_x, page_h - margin - 58, sidebar_w)
    c.setFillColor(colors.HexColor("#1F4E79"))
    c.setFont(FONT_BOLD, 20)
    c.drawString(main_x, page_h - margin, cv["name"])
    c.setFont(FONT_BOLD, 10)
    c.setFillColor(colors.black)
    c.drawString(main_x, page_h - margin - 16, cv["title"])
    c.setStrokeColor(colors.HexColor("#1F4E79"))
    c.line(main_x, page_h - margin - 25, sidebar_x - gap, page_h - margin - 25)

    y = page_h - margin - 42

    def ensure_local(needed):
        nonlocal y
        if y - needed < margin:
            c.showPage()
            draw_sidebar(c, cv, sidebar_x, page_h - margin, sidebar_w)
            y = page_h - margin

    def heading_local(text):
        nonlocal y
        ensure_local(23)
        y -= 4
        c.setFillColor(colors.HexColor("#1F4E79"))
        c.setFont(FONT_BOLD, 9.5)
        c.drawString(main_x, y, text.upper())
        y -= 4
        c.setStrokeColor(colors.HexColor("#1F4E79"))
        c.line(main_x, y, main_x + main_w, y)
        y -= 10

    def write_local(text, size=8.3, font=FONT_REGULAR, leading=9.8):
        nonlocal y
        lines = split_text(text, font, size, main_w)
        ensure_local(len(lines) * leading + 2)
        c.setFont(font, size)
        c.setFillColor(colors.black)
        for line in lines:
            c.drawString(main_x, y, line)
            y -= leading
        y -= 1

    def bullet_local(text):
        nonlocal y
        lines = split_text(text, FONT_REGULAR, 8.1, main_w - 10)
        ensure_local(len(lines) * 9.4 + 3)
        c.setFont(FONT_REGULAR, 8.1)
        c.setFillColor(colors.black)
        c.drawString(main_x, y, "-")
        c.drawString(main_x + 10, y, lines[0])
        y -= 9.4
        for line in lines[1:]:
            c.drawString(main_x + 10, y, line)
            y -= 9.4
        y -= 1

    heading_local(cv["profile_title"])
    write_local(cv["profile"])

    for section in cv["sections"]:
        if section["title"] in ("Key Skills", "Kluczowe umiejętności"):
            continue
        heading_local(section["title"])
        if "subtitle" in section:
            write_local(section["subtitle"], size=8.4, font=FONT_BOLD, leading=10)
        for item in section.get("items", []):
            bullet_local(item)
        for job in section.get("jobs", []):
            write_local(f"{job['role']} | {job['company']} | {job['date']}", size=8.4, font=FONT_BOLD, leading=10)
            for item in job["items"]:
                bullet_local(item)

    heading_local("Consent" if cv is CV_EN else "Zgoda")
    write_local(cv["consent"], size=6.8, leading=8)
    c.save()
    return path


def write_recruiter_messages() -> Path:
    path = OUT_DIR / "recruiter_messages_softserve.md"
    path.write_text(
        """# Wiadomości do rekrutera - SoftServe Junior Robotics Engineer

## Polish

Dzień dobry,

przesyłam aplikację na stanowisko Junior Robotics Engineer. Jako student Mechaniki i Budowy Maszyn oraz Technik Programista łączę praktyczne projektowanie mechaniczne z programowaniem i symulacjami robotycznymi. Mam komercyjne doświadczenie w tworzeniu symulacji robotycznych i cyfrowych bliźniaków w Visual Components, w tym pracy ze ścieżkami ruchu robotów, logiką procesu, wykrywaniem kolizji i optymalizacją layoutu.

W projektach własnych zaprojektowałem i zbudowałem prototyp ramienia robota SCARA, obejmujący CAD, druk 3D, dobór napędów i sterowników, elektronikę oraz kod sterujący w Pythonie/C++. Chętnie rozwinę te kompetencje w zespole Robotics Group przy projektach z zakresu symulacji, prototypowania i integracji systemów robotycznych.

Chętnie opowiem więcej o moich projektach podczas rozmowy.

Pozdrawiam,
Maciej Tkacz

## English

Hello,

I would like to apply for the Junior Robotics Engineer position. As a Mechanical Engineering student and certified Software Technician, I combine hands-on mechanical design with programming and robotics simulation. I have commercial experience creating robotic simulations and digital twins in Visual Components, including robot paths, process logic, collision checking and layout optimization.

In my personal projects, I designed and built a SCARA robotic arm prototype covering CAD, 3D printing, actuator and driver selection, electronics, and Python/C++ control software. I would be glad to further develop these skills within the Robotics Group, especially in simulation, prototyping and hardware-software integration of robotic systems.

I would be happy to discuss my projects and motivation in more detail.

Best regards,
Maciej Tkacz
""",
        encoding="utf-8",
    )
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    created = []
    for cv in (CV_EN, CV_PL):
        created.append(build_ats_docx(cv))
        created.append(build_visual_docx(cv))
        created.append(build_ats_pdf(cv))
        created.append(build_visual_pdf(cv))
    created.append(write_recruiter_messages())
    for path in created:
        print(path.relative_to(OUT_DIR.parents[1]))


if __name__ == "__main__":
    main()
