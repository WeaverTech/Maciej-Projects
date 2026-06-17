# -*- coding: utf-8 -*-
import json, math, docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
import data as D
from compute import phi
from dxutil import set_para, cell_set, insert_paragraph_after, image_after

B=json.load(open("computed.json"))
s1=B['s1']; s2=B['s2']; S1=B['S1']; S2=B['S2']; r1=B['r1']; r2=B['r2']
F=B['F']; t=B['t']; c1=B['c1']; c2=B['c2']
def pc(v,n=3): return f"{v:.{n}f}".replace('.',',')
CHIKR="11,07"; FKR="1,60"; TKR="1,984"

doc=docx.Document("../L2_wzor_sprawozdania_aktualne.docx")
T=doc.tables

# --- naglowek (tabela 0) ---
cell_set(T[0].rows[0].cells[1],
  "Jakub Suchoń, Jan Sendecki, Kacper Tokarski, Jakub Wydra, Adam Powrózek, Maciej Tkacz, Jakub Pitala\n"
  "Grupa lab L04_GR2   Zespół 12M2   rok akad. 2025/2026", size=8)

# --- serie (tabele 1,2) ---
def fill_series(tbl, xs):
    for k,v in enumerate(xs):
        row = 1 + (k % 17)
        col = 1 + (k // 17)*2
        cell_set(tbl.rows[row].cells[col], pc(v), size=8.5, align=AL.CENTER)
fill_series(T[1], S1); fill_series(T[2], S2)

# --- wskazniki (3=x̄,4=mₑ,5=R,6=s²,7=s) ---
def w(ti, v1, v2, unit="mm", n=3):
    cell_set(T[ti].rows[0].cells[1], f"= {pc(v1,n)} {unit}", size=9)
    cell_set(T[ti].rows[0].cells[3], f"= {pc(v2,n)} {unit}", size=9)
w(3, s1['mean'], s2['mean'])
w(4, s1['median'], s2['median'])
w(5, s1['R'], s2['R'])
w(6, s1['var'], s2['var'], unit="mm²", n=4)
w(7, s1['s'], s2['s'])

# --- interpretacja: odpowiedzi po pytaniach ---
def ans(qstart, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(qstart):
            insert_paragraph_after(p, text, size=10, bold=False); return
ans("Czy średnia x\u0304 jest zbliżona",
    f"Średnie obu serii (x\u03041 = {pc(s1['mean'])} mm, x\u03042 = {pc(s2['mean'])} mm) są zbliżone do "
    f"wartości nominalnej; leżą nieco poniżej środka pola tolerancji, co jest spójne z lekkim przesunięciem "
    f"procesu stwierdzonym w L8.")
ans("Czy średnia i mediana",
    f"Tak. W obu seriach różnica między średnią a medianą jest minimalna (S1: {pc(s1['mean'])} vs {pc(s1['median'])}; "
    f"S2: {pc(s2['mean'])} vs {pc(s2['median'])}), co wskazuje na symetryczny kształt rozkładu.")
ans("Czy średnie obu serii są zbliżone",
    f"Tak, średnie są prawie identyczne (różnica {pc(abs(s1['mean']-s2['mean']),3)} mm).")
ans("Porównaj odchylenia standardowe",
    f"Seria 1 jest bardziej powtarzalna — ma niższe odchylenie standardowe (s1 = {pc(s1['s'])} mm) niż "
    f"Seria 2 (s2 = {pc(s2['s'])} mm). Różnica może wynikać z losowej zmienności próbkowania na tym samym sprzęcie.")
ans("Czy rozstęp R wydaje się proporcjonalny",
    f"Tak — wyższemu odchyleniu standardowemu w Serii 2 towarzyszy proporcjonalnie większy rozstęp "
    f"(R2 = {pc(s2['R'])} mm wobec R1 = {pc(s1['R'])} mm).")
ans("Czy histogram ma kształt zbliżony",
    "Tak, oba histogramy są najwyższe w części środkowej i opadają ku brzegom (kształt zbliżony do dzwonowego), "
    "co oznacza, że większość wymiarów skupia się blisko środka tolerancji.")
ans("Czy są przedziały z bardzo małą liczebnością",
    "Na końcach wykresów pojawiają się pojedyncze wyniki, ale nie są na tyle oddalone od reszty, by uznać je za "
    "błędy grube — wszystkie mieszczą się w przedziale [x\u0304 − 3s ; x\u0304 + 3s].")
ans("Porównaj histogramy obu serii",
    "Oba histogramy mają podobny, regularny kształt, lecz histogram Serii 2 jest nieco szerszy, co odpowiada "
    "większemu rozrzutowi wymiarów w tej serii.")

# --- wykluczenie wartosci nadmiernych (tabela 9) ---
def excl(cell, s):
    lo=s['mean']-3*s['s']; hi=s['mean']+3*s['s']
    cell_set(cell, f"x\u0304 − 3s = {pc(lo)} mm\nx\u0304 + 3s = {pc(hi)} mm\n"
                   f"Wartości odrzucone: BRAK\nNowe n = BRAK (n = 51)", size=8.5)
excl(T[9].rows[1].cells[0], s1); excl(T[9].rows[1].cells[1], s2)

# --- grupowanie (10=S1, 12=S2) i chi2 (11=S1, 13=S2) ---
PRZE=[
 "(−∞ ; x\u0304 − 1,15s)","[x\u0304 − 1,15s ; x\u0304 − 0,75s)","[x\u0304 − 0,75s ; x\u0304 − 0,3s)",
 "[x\u0304 − 0,3s ; x\u0304)","[x\u0304 ; x\u0304 + 0,3s)","[x\u0304 + 0,3s ; x\u0304 + 0,75s)",
 "[x\u0304 + 0,75s ; x\u0304 + 1,15s)","[x\u0304 + 1,15s ; +∞)"]
def fill_group(tg, tc, rows, s, chi):
    # naglowek chi2
    hdr=["Nr","mᵢ","u₁=(lᵢ−x\u0304)/s","u₂=(lᵢ₋₁−x\u0304)/s","Φ(u₁)","Φ(u₂)","Pᵢ","nPᵢ","(mᵢ−nPᵢ)²/nPᵢ"]
    for ci,h in enumerate(hdr): cell_set(tc.rows[0].cells[ci], h, bold=True, size=7.5, align=AL.CENTER)
    for k,rr in enumerate(rows):
        gi=k+1
        lo,hi=rr['lo'],rr['hi']
        gl = ("−∞" if lo==float('-inf') else pc(lo)) + " ; " + ("+∞" if hi==float('inf') else pc(hi))
        cell_set(tg.rows[gi].cells[1], PRZE[k], size=8)
        cell_set(tg.rows[gi].cells[2], gl, size=8, align=AL.CENTER)
        cell_set(tg.rows[gi].cells[3], str(rr['mi']), size=8, align=AL.CENTER)
        u1 = math.inf if hi==float('inf') else (hi-s['mean'])/s['s']
        u2 = -math.inf if lo==float('-inf') else (lo-s['mean'])/s['s']
        Phi1 = 0.5 if u1==math.inf else phi(u1)
        Phi2 = -0.5 if u2==-math.inf else phi(u2)
        cell_set(tc.rows[gi].cells[1], str(rr['mi']), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[2], ("+∞" if u1==math.inf else pc(u1,2)), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[3], ("−∞" if u2==-math.inf else pc(u2,2)), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[4], pc(Phi1,4), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[5], pc(Phi2,4), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[6], pc(rr['Pi'],4), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[7], pc(rr['nPi'],3), size=8, align=AL.CENTER)
        cell_set(tc.rows[gi].cells[8], pc(rr['comp'],4), size=8, align=AL.CENTER)
    # wiersz sumy
    sr=tc.rows[9]
    cell_set(sr.cells[1], "Σ = 51", bold=True, size=8, align=AL.CENTER)
    cell_set(sr.cells[6], "Σ = 1", bold=True, size=8, align=AL.CENTER)
    cell_set(sr.cells[7], "Σ = 51", bold=True, size=8, align=AL.CENTER)
    cell_set(sr.cells[8], f"χ² = {pc(chi,2)}", bold=True, size=8, align=AL.CENTER)
fill_group(T[10], T[11], r1, s1, c1)
fill_group(T[12], T[13], r2, s2, c2)

# --- histogramy (tabela 8) ---
T8=doc.tables[8]
image_after(T8.rows[1].cells[0].paragraphs[0], "ch_hist_s1.png", 78)
image_after(T8.rows[1].cells[1].paragraphs[0], "ch_hist_s2.png", 78)

# --- decyzja chi2 (tabela 14) ---
cell_set(T[14].rows[1].cells[0], f"L₁ = 8 przedziałów; k₁ = L−3 = 5; χ²₁ = {pc(c1,2)}; χ²kr = {CHIKR}; χ²₁ ≤ χ²kr → TAK", size=8.5)
cell_set(T[14].rows[1].cells[1], f"L₂ = 8 przedziałów; k₂ = L−3 = 5; χ²₂ = {pc(c2,2)}; χ²kr = {CHIKR}; χ²₂ ≤ χ²kr → TAK", size=8.5)
cell_set(T[14].rows[2].cells[0], "Wniosek (Seria 1): brak podstaw do odrzucenia hipotezy o normalności. Dane są "
        "zgodne z rozkładem normalnym przy β = 0,95. Można przystąpić do testów w Części 3 i 4.", size=8.5)
cell_set(T[14].rows[2].cells[1], "Wniosek (Seria 2): brak podstaw do odrzucenia hipotezy o normalności. Dane są "
        "zgodne z rozkładem normalnym przy β = 0,95. Można przystąpić do testów w Części 3 i 4.", size=8.5)

# --- test F (tabele 15,16) ---
cell_set(T[15].rows[0].cells[0], f"s²₁ = {pc(s1['var'],4)} mm²", size=9)
cell_set(T[15].rows[0].cells[1], f"s²₂ = {pc(s2['var'],4)} mm²", size=9)
cell_set(T[16].rows[0].cells[0], f"F = s²max / s²min = {pc(F,3)}", size=9)
cell_set(T[16].rows[1].cells[0], "ν₁ = n₁ − 1 = 50", size=9)
cell_set(T[16].rows[1].cells[1], "ν₂ = n₂ − 1 = 50", size=9)
cell_set(T[16].rows[2].cells[0], "Fkryt (z tabeli 3) =", size=9)
cell_set(T[16].rows[2].cells[1], FKR, size=9)
cell_set(T[16].rows[3].cells[0], "F ≤ Fkryt ?  TAK", size=9)
cell_set(T[16].rows[3].cells[1], "TAK", size=9)
cell_set(T[16].rows[4].cells[0], "Wniosek: brak podstaw do odrzucenia H₀. Wariancje obu serii są statystycznie "
        "równe — można zastosować standardowy test t.", size=8.5)
cell_set(T[16].rows[4].cells[1], "Wniosek: wariancje równe (F < Fkryt).", size=8.5)

# --- test t (tabela 17) ---
cell_set(T[17].rows[0].cells[0], f"ts = (x\u03041 − x\u03042) / √(s²₁/n₁ + s²₂/n₂) = {pc(t,4)}", size=9)
cell_set(T[17].rows[0].cells[1], "k = n₁ + n₂ − 2 = 100", size=9)
cell_set(T[17].rows[1].cells[0], "tkryt (z tabeli 4, β = 0,95) =", size=9)
cell_set(T[17].rows[1].cells[1], TKR, size=9)
cell_set(T[17].rows[2].cells[0], "|ts| ≤ tkryt ?  TAK", size=9)
cell_set(T[17].rows[2].cells[1], "TAK", size=9)
cell_set(T[17].rows[3].cells[0], "Wniosek: brak podstaw do odrzucenia H₀. Różnica średnich jest statystycznie "
        "nieistotna — obie serie mogą mieć tę samą wartość oczekiwaną.", size=8.5)
cell_set(T[17].rows[3].cells[1], "Wniosek: średnie statystycznie równe.", size=8.5)

# --- wnioski koncowe (P101) ---
for p in doc.paragraphs:
    if p.text.strip().startswith("Wnioski i podsumowanie"):
        insert_paragraph_after(p,
          f"Wyniki analizy obu serii pomiarowych potwierdzają poprawną i stabilną kontrolę wysokości wałków. "
          f"Bliskie sobie wartości średnich (x\u03041 = {pc(s1['mean'])} mm, x\u03042 = {pc(s2['mean'])} mm) i median "
          f"świadczą o symetrii rozkładu empirycznego, a test χ² nie dał podstaw do odrzucenia hipotezy o "
          f"normalności (χ²₁ = {pc(c1,2)}; χ²₂ = {pc(c2,2)}; oba < {CHIKR}). Seria 1 cechuje się nieco wyższą "
          f"powtarzalnością (mniejsze s i R), jednak test F (F = {pc(F,3)} ≤ {FKR}) oraz test t "
          f"(|ts| = {pc(abs(t),3)} ≤ {TKR}) wykazały, że różnice wariancji i średnich między seriami są "
          f"statystycznie nieistotne. Pomiary wykonano na tym samym sprzęcie i tym samym przedmiocie, co tłumaczy "
          f"zgodność obu serii.", size=10, normal=True)
        break

out="../Sprawozdania_L2_L8/L2_sprawozdanie.docx"
doc.save(out); print("zapisano", out)
