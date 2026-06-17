# -*- coding: utf-8 -*-
import copy
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, RGBColor

def set_para(p, text, bold=None, size=None):
    """Ustawia tekst akapitu zachowujac styl pierwszego run-a."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
        run = p.runs[0]
    else:
        run = p.add_run(text)
    if bold is not None: run.bold = bold
    if size is not None: run.font.size = Pt(size)
    return p

def cell_set(cell, text, bold=False, size=8.5, align=None):
    """Wpisuje tekst do komorki (jedna linijka), zachowujac prosty styl."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return cell

def insert_paragraph_after(paragraph, text="", style=None, bold=False, size=None, normal=False):
    new_p = copy.deepcopy(paragraph._p)
    # wyczysc zawartosc skopiowanego akapitu
    for child in list(new_p):
        if child.tag == qn('w:r') or child.tag == qn('w:hyperlink'):
            new_p.remove(child)
    if normal:
        # usun numeracje/listy i odstep, ustaw zwykly akapit
        pPr = new_p.find(qn('w:pPr'))
        if pPr is not None:
            for tag in ('w:numPr', 'w:pStyle', 'w:outlineLvl'):
                el = pPr.find(qn(tag))
                if el is not None: pPr.remove(el)
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try: np.style = style
        except Exception: pass
    if text:
        r = np.add_run(text); r.bold = bold
        if size: r.font.size = Pt(size)
    return np

def clone_row(table, src_idx):
    src_tr = table.rows[src_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    src_tr.addnext(new_tr)
    return table.rows[src_idx+1]

def image_after(paragraph, path, width_mm=150, center=True):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    np = insert_paragraph_after(paragraph, "")
    if center:
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.add_run().add_picture(path, width=Mm(width_mm))
    return np
