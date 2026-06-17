# -*- coding: utf-8 -*-
import json, docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
import data as D
from dxutil import set_para, cell_set, insert_paragraph_after, clone_row, image_after

B=json.load(open("computed.json")); L=B['L']; L8=B['L8']
def pc(v, n=3): return f"{v:.{n}f}".replace('.', ',')   # liczba z przecinkiem dz.

doc=docx.Document("../Lab8_wzor_spr.docx")

# ---- naglowek ----
t0=doc.tables[0]
cell_set(t0.rows[0].cells[1], "Jakub Suchoń, Jan Sendecki, Kacper Tokarski, Jakub Wydra", size=8)
cell_set(t0.rows[0].cells[2], "Adam Powrózek, Maciej Tkacz, Jakub Pitala", size=8)
cell_set(t0.rows[1].cells[1], "grupa: "+D.GRUPA, size=9)
cell_set(t0.rows[1].cells[2], "rok akad.: "+D.ROK, size=9)
cell_set(t0.rows[2].cells[1], "grupa lab: "+D.GRUPA_LAB, size=9)
cell_set(t0.rows[2].cells[2], "zespół: "+D.ZESPOL, size=9)

# ---- tabela danych ----
t=doc.tables[1]
means=[sum(s)/5 for s in L8]; Rs=[max(s)-min(s) for s in L8]
def fill_data_row(row, i):
    smp=L8[i]; m=means[i]; mx=max(smp); mn=min(smp); R=Rs[i]
    uw=[]
    if m>L['UCLx'] or m<L['LCLx']: uw.append("x̄ poza granicą")
    if R>L['UCLr']: uw.append("R > UCL")
    vals=[str(i+1)]+[pc(v) for v in smp]+[pc(m),pc(mx),pc(mn),pc(R),"; ".join(uw)]
    for ci,v in enumerate(vals):
        cell_set(row.cells[ci], v, size=8, align=AL.CENTER)
fill_data_row(t.rows[3], 0)
for i in range(1,25):
    nr=clone_row(t, 3+i-1)
    fill_data_row(nr, i)
tot=t.rows[-1]
cell_set(tot.cells[0], "Podsumowanie", bold=True, size=8, align=AL.CENTER)
cell_set(tot.cells[6], "x̿ = "+pc(L['xbb']), bold=True, size=8, align=AL.CENTER)
cell_set(tot.cells[9], "R̄ = "+pc(L['Rb']), bold=True, size=8, align=AL.CENTER)

def by_text(s, starts=False):
    for p in doc.paragraphs:
        if (p.text.strip().startswith(s) if starts else p.text.strip()==s): return p
    return None

# ---- teoria ----
p=by_text("Rodzaje kart kontrolnych")
for txt in reversed([
 "Karty kontrolne (karty Shewharta) to podstawowe narzędzie statystycznego sterowania procesem (SPC). "
 "Służą do bieżącego monitorowania stabilności procesu oraz wczesnego wykrywania rozregulowań. Dzieli się je na:",
 "A. Karty dla cech mierzalnych: karta x̄ (średniej) – nadzoruje wycentrowanie procesu; karta R (rozstępu) – "
 "kontroluje zmienność w próbce; karta S (odchylenia standardowego) – dla większych próbek; karta I-MR – gdy n = 1.",
 "B. Karty dla cech atrybutowych: karta p – frakcja wyrobów wadliwych; karta np – liczba wadliwych przy stałej "
 "liczności; karta c – liczba wad na jednostce; karta u – liczba wad na jednostkę przy zmiennej próbce.",
]): insert_paragraph_after(p, txt, size=10)

p=by_text("Charakterystyka parametrów Cp i Cpk")
for txt in reversed([
 "Wskaźnik zdolności potencjalnej Cp = T/(6·Sproc) = (GWG − DWG)/(6·Sproc) określa teoretyczną zdolność procesu, "
 "bez uwzględnienia wycentrowania.",
 "Wskaźnik zdolności rzeczywistej Cpk = min[(GWG − x̄)/(3·Sproc); (x̄ − DWG)/(3·Sproc)] uwzględnia przesunięcie "
 "średniej względem środka pola tolerancji. Gdy Cp = Cpk – proces wycentrowany; Cp < 1 – proces niezdolny; "
 "Cp ≠ Cpk – proces źle ustawiony.",
 "Oznaczenia: T – tolerancja, GWG/DWG – górny/dolny wymiar graniczny, x̄ – średnia procesu, Sproc – eksperymentalne "
 "odchylenie standardowe procesu (Sproc = R̄/d2).",
]): insert_paragraph_after(p, txt, size=10)

# ---- linie kontrolne (kolejnosc dok.: CLx,UCLx,LCLx,CLr,UCLr,LCLr) ----
PP=doc.paragraphs
order=[i for i,p in enumerate(PP) if p.text.strip() in ("CL = ?","UCL= ?","LCL = ?")]
vals=[L['xbb'],L['UCLx'],L['LCLx'],L['Rb'],L['UCLr'],L['LCLr']]
for idx,val in zip(order,vals):
    set_para(PP[idx], PP[idx].text.replace("?", pc(val,4)))

# ---- Cp/Cpk ----
for p in doc.paragraphs:
    if p.text.strip()=="Cp=?": set_para(p, "Cp = "+pc(L['Cp']))
    if p.text.strip()=="Cpk=?": set_para(p, "Cpk = "+pc(L['Cpk']))
pg=by_text("Wartości GWG i DWG", starts=True)
insert_paragraph_after(pg,
    f"Przyjęto (wg prowadzącego): GWG = {pc(D.GWG)} mm, DWG = {pc(D.DWG)} mm, T = {pc(D.GWG-D.DWG)} mm. "
    f"Sproc = R̄/d2 = {pc(L['Rb'])}/2,326 = {pc(L['Sproc'],4)} mm (d2 = 2,326 dla n = 5).", size=10, normal=True)

# ---- wykresy ----
p=by_text("Opracować kartę kontrolną", starts=True)
image_after(p, "ch_hist_means.png", 110)
image_after(p, "ch_r.png", 165)
image_after(p, "ch_x.png", 165)

# ---- wnioski ----
maxR=max(Rs)
pw=by_text("INTERPRETACJA WYNIKÓW, PODSUMOWANIE I WNIOSKI")
wn=(f"Analiza kart kontrolnych x̄ oraz R wykazuje, że proces nie jest w pełni stabilny statystycznie. "
    f"Na karcie R próbki nr 1 oraz 21 przekraczają górną granicę kontrolną (UCL_R = {pc(L['UCLr'])}), "
    f"a na karcie x̄ próbka nr 1 leży poniżej dolnej granicy (LCL = {pc(L['LCLx'])}); świadczy to o dużej "
    f"zmienności wewnątrz tych próbek (średni rozstęp R̄ = {pc(L['Rb'])} mm, maksymalnie do {pc(maxR)} mm). "
    f"Wskaźnik zdolności Cp = {pc(L['Cp'])} jest mniejszy od 1, co oznacza, że proces nie jest zdolny do "
    f"spełnienia wymagań tolerancyjnych nawet przy idealnym wycentrowaniu. Ponieważ Cp ≠ Cpk "
    f"(Cpk = {pc(L['Cpk'])}), proces jest źle ustawiony – średnia ogólna jest przesunięta w stronę dolnego "
    f"wymiaru granicznego. W konsekwencji badany proces nie spełnia wymagań jakościowych wyrobu i wymaga "
    f"optymalizacji: wycentrowania ustawienia maszyny oraz zmniejszenia rozrzutu obróbki.")
insert_paragraph_after(pw, wn, size=10, normal=True)

out="../Sprawozdania_L2_L8/L8_sprawozdanie.docx"
doc.save(out); print("zapisano", out)
